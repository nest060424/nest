from django.shortcuts import render, redirect

# Класс HttpResponse из пакета django.http, который позволяет отправить текстовое содержимое.
from django.http import HttpResponse, HttpResponseNotFound
# Конструктор принимает один обязательный аргумент – путь для перенаправления. Это может быть полный URL (например, 'https://www.yahoo.com/search/') или абсолютный путь без домена (например, '/search/').
from django.http import HttpResponseRedirect

from django.urls import reverse

from django.contrib.auth.decorators import login_required
from django.contrib.auth.decorators import user_passes_test

# Подключение моделей
from .models import Category, Department, Position, Employee, Person, Status, Statement, Kind, Queue, History, News 
# Подключение форм
from .forms import CategoryForm, DepartmentForm, PositionForm, EmployeeForm, PersonForm, StatusForm, StatementCreateForm, StatementCheckForm, StatementSubscribeForm, StatementEditForm, QueueEditForm, KindForm, SignUpForm, NewsForm
#from .forms import SignUpForm

import datetime

import xlwt
from io import BytesIO

from django.db import models

import sys

#from django.utils.translation import ugettext as _
from django.utils.translation import gettext_lazy as _

from django.utils.decorators import method_decorator
from django.views.generic import UpdateView
from django.contrib.auth.models import User, Group
from django.urls import reverse_lazy

from django.contrib.auth import login as auth_login

from django.shortcuts import get_object_or_404

import random

from docx import Document
from docx.shared import Inches

# Групповые ограничения
def group_required(*group_names):
    """Requires user membership in at least one of the groups passed in."""
    def in_groups(u):
        if u.is_authenticated:
            if bool(u.groups.filter(name__in=group_names)) | u.is_superuser:
                return True
        return False
    return user_passes_test(in_groups, login_url='403')

# Стартовая страница
#@login_required 
def index(request):
    return render(request, "index.html")

# Страница Контакты
def contact(request):
    return render(request, "contact.html")

# Страница Полезных ссылок
def link(request):
    return render(request, "link.html")

# Страница Полезных ссылок
def faq(request):
    return render(request, "faq.html")

# Страница Отчеты
@login_required
@group_required("Managers", "Managements")
def report(request):
    # Заявление
    statement = Statement.objects.all().order_by('-dates')
    # Очередь
    queue = Queue.objects.filter(dateq2__isnull=True).order_by('dateq')   
    # Выбрать все статусы для фильтра по статусам
    status = Status.objects.all().order_by('title')
    # Выбрать все категории для фильтра по категориям
    category = Category.objects.all().order_by('title')
    if request.method == "POST":
        # Определить какая кнопка нажата
        if 'btnStatement' in request.POST:
            print('btnStatement')
            # Выбранный статус
            selected_item_id = request.POST.get('item_status_id')
            print(selected_item_id)
            statement = Statement.objects.filter(status_id=selected_item_id).order_by('-dates')        
        if 'btnQueue' in request.POST:
            # Выбрання категория
            print('btnQueue')
            selected_item_id = request.POST.get('item_category_id')
            print(selected_item_id)
            #queue = Queue.objects.filter(statement_id=selected_item_id).order_by('dateq')
            statement_query = Statement.objects.filter(category_id=selected_item_id).only('id').all()
            print(statement_query)
            queue = Queue.objects.filter(statement_id__in=statement_query)
    return render(request, "report.html", {"statement": statement, "signature": request.user.employee.signature, "status": status, "queue": queue, "category": category,})

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def category_index(request):
    category = Category.objects.all().order_by('title')
    return render(request, "category/index.html", {"category": category})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def category_create(request):
    if request.method == "POST":
        category = Category()        
        category.title = request.POST.get("title")
        category.save()
        return HttpResponseRedirect(reverse('category_index'))
    else:        
        categoryform = CategoryForm(request.FILES)
        return render(request, "category/create.html", {"form": categoryform})

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def category_edit(request, id):
    try:
        category = Category.objects.get(id=id) 
        if request.method == "POST":
            category.title = request.POST.get("title")
            category.save()
            return HttpResponseRedirect(reverse('category_index'))
        else:
            # Загрузка начальных данных
            categoryform = CategoryForm(initial={'title': category.title, })
            return render(request, "category/edit.html", {"form": categoryform})
    except Category.DoesNotExist:
        return HttpResponseNotFound("<h2>Category not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def category_delete(request, id):
    try:
        category = Category.objects.get(id=id)
        category.delete()
        return HttpResponseRedirect(reverse('category_index'))
    except Category.DoesNotExist:
        return HttpResponseNotFound("<h2>Category not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def category_read(request, id):
    try:
        category = Category.objects.get(id=id) 
        return render(request, "category/read.html", {"category": category})
    except Category.DoesNotExist:
        return HttpResponseNotFound("<h2>Category not found</h2>")    

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def department_index(request):
    department = Department.objects.all().order_by('title')
    return render(request, "department/index.html", {"department": department})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def department_create(request):
    if request.method == "POST":
        department = Department()        
        department.title = request.POST.get("title")
        department.save()
        return HttpResponseRedirect(reverse('department_index'))
    else:        
        departmentform = DepartmentForm(request.FILES)
        return render(request, "department/create.html", {"form": departmentform})

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def department_edit(request, id):
    try:
        department = Department.objects.get(id=id) 
        if request.method == "POST":
            department.title = request.POST.get("title")
            department.save()
            return HttpResponseRedirect(reverse('department_index'))
        else:
            # Загрузка начальных данных
            departmentform = DepartmentForm(initial={'title': department.title, })
            return render(request, "department/edit.html", {"form": departmentform})
    except Department.DoesNotExist:
        return HttpResponseNotFound("<h2>Department not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def department_delete(request, id):
    try:
        department = Department.objects.get(id=id)
        department.delete()
        return HttpResponseRedirect(reverse('department_index'))
    except Department.DoesNotExist:
        return HttpResponseNotFound("<h2>Department not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def department_read(request, id):
    try:
        department = Department.objects.get(id=id) 
        return render(request, "department/read.html", {"department": department})
    except Department.DoesNotExist:
        return HttpResponseNotFound("<h2>Department not found</h2>") 

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def position_index(request):
    position = Position.objects.all().order_by('title')
    return render(request, "position/index.html", {"position": position})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def position_create(request):
    if request.method == "POST":
        position = Position()        
        position.title = request.POST.get("title")
        position.save()
        return HttpResponseRedirect(reverse('position_index'))
    else:        
        positionform = PositionForm(request.FILES)
        return render(request, "position/create.html", {"form": positionform})

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def position_edit(request, id):
    try:
        position = Position.objects.get(id=id) 
        if request.method == "POST":
            position.title = request.POST.get("title")
            position.save()
            return HttpResponseRedirect(reverse('position_index'))
        else:
            # Загрузка начальных данных
            positionform = PositionForm(initial={'title': position.title, })
            return render(request, "position/edit.html", {"form": positionform})
    except Position.DoesNotExist:
        return HttpResponseNotFound("<h2>Position not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def position_delete(request, id):
    try:
        position = Position.objects.get(id=id)
        position.delete()
        return HttpResponseRedirect(reverse('position_index'))
    except Position.DoesNotExist:
        return HttpResponseNotFound("<h2>Position not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def position_read(request, id):
    try:
        position = Position.objects.get(id=id) 
        return render(request, "position/read.html", {"position": position})
    except Position.DoesNotExist:
        return HttpResponseNotFound("<h2>Position not found</h2>")

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def employee_index(request):
    employee = Employee.objects.all().order_by('surname', 'name', 'patronymic')
    return render(request, "employee/index.html", {"employee": employee})

# Функция edit выполняет редактирование объекта.
@login_required
def employee_edit(request, id):
    try:
        employee = Employee.objects.get(id=id)
        if request.method == "POST":
            employee.surname = request.POST.get("surname")
            employee.name = request.POST.get("name")
            employee.patronymic = request.POST.get("patronymic")
            employee.department = Department.objects.filter(id=request.POST.get("department")).first()
            employee.position = Position.objects.filter(id=request.POST.get("position")).first()
            employee.signature = request.POST.get("signature")
            employee.save()
            return HttpResponseRedirect(reverse('employee_index'))
        else:
            # Загрузка начальных данных
            employeeform = EmployeeForm(initial={'surname': employee.surname, 'name': employee.name, 'patronymic': employee.patronymic,  'department': employee.department, 'position': employee.position, 'signature': employee.signature,})
            return render(request, "employee/edit.html", {"form": employeeform})
    except Department.DoesNotExist:
        return HttpResponseNotFound("<h2>Employee not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
@group_required("Managers")
def employee_read(request, id):
    try:
        employee = Employee.objects.get(id=id) 
        return render(request, "employee/read.html", {"employee": employee})
    except Groups.DoesNotExist:
        return HttpResponseNotFound("<h2>Employee not found</h2>")

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def person_index(request):
    person = Person.objects.all().order_by('surname', 'name', 'patronymic')
    return render(request, "person/index.html", {"person": person})

# Функция edit выполняет редактирование объекта.
@login_required
#def person_edit(request, id):
def person_edit(request):
    try:
        #person = Person.objects.get(id=id)
        print(request.user.id)
        person = Person.objects.get(user_id=request.user.id)        
        if request.method == "POST":
            person.sex = request.POST.get("sex")
            person.iin = request.POST.get("iin")
            person.birthday = request.POST.get("birthday")
            person.surname = request.POST.get("surname")
            person.name = request.POST.get("name")
            person.patronymic = request.POST.get("patronymic")
            person.save()
            return HttpResponseRedirect(reverse('index'))
        else:
            # Загрузка начальных данных
            personform = PersonForm(initial={'sex': person.sex, 'iin': person.iin, 'birthday': person.birthday.strftime('%Y-%m-%d'), 'surname': person.surname, 'name': person.name, 'patronymic': person.patronymic, })
            return render(request, "person/edit.html", {"form": personform})
    except Department.DoesNotExist:
        return HttpResponseNotFound("<h2>Person not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
@group_required("Managers")
def person_read(request, id):
    try:
        person = Person.objects.get(id=id) 
        return render(request, "person/read.html", {"person": person})
    except Groups.DoesNotExist:
        return HttpResponseNotFound("<h2>Person not found</h2>")

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def status_index(request):
    status = Status.objects.all().order_by('title')
    return render(request, "status/index.html", {"status": status})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def status_create(request):
    if request.method == "POST":
        status = Status()        
        status.title = request.POST.get("title")
        status.save()
        return HttpResponseRedirect(reverse('status_index'))
    else:        
        statusform = StatusForm(request.FILES)
        return render(request, "status/create.html", {"form": statusform})

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def status_edit(request, id):
    try:
        status = Status.objects.get(id=id) 
        if request.method == "POST":
            status.title = request.POST.get("title")
            status.save()
            return HttpResponseRedirect(reverse('status_index'))
        else:
            # Загрузка начальных данных
            statusform = StatusForm(initial={'title': status.title, })
            return render(request, "status/edit.html", {"form": statusform})
    except Status.DoesNotExist:
        return HttpResponseNotFound("<h2>Status not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def status_delete(request, id):
    try:
        status = Status.objects.get(id=id)
        status.delete()
        return HttpResponseRedirect(reverse('status_index'))
    except Status.DoesNotExist:
        return HttpResponseNotFound("<h2>Status not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def status_read(request, id):
    try:
        status = Status.objects.get(id=id) 
        return render(request, "status/read.html", {"status": status})
    except Status.DoesNotExist:
        return HttpResponseNotFound("<h2>Status not found</h2>")    

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers", "Commissions")
def statement_index(request):
    #statement = Statement.objects.all().order_by('dates')
    #print('signature', request.user.employee.signature) 
    #return render(request, "statement/index.html", {"statement": statement, "signature": request.user.employee.signature})
    # Выбрать все статусы для фильтра по статусам
    status = Status.objects.all().order_by('title')
    if request.method == "POST":
        # Выбранный статус
        selected_item_id = request.POST.get('item_id')
        statement = Statement.objects.filter(status_id=selected_item_id).order_by('-dates')        
    else:
        statement = Statement.objects.all().order_by('-dates')
    return render(request, "statement/index.html", {"statement": statement, "signature": request.user.employee.signature, "status": status,})

# Список 
@login_required
def statement_list(request):
    #statement = Statement.objects.all().order_by('dates')
    print(request.user.person.id)
    statement = Statement.objects.filter(person_id=request.user.person.id).order_by('-dates')
    # Если заявлений пока нет то и показывать нечего 
    if statement.count()==0:
        return render(request, "statement/list.html", {"statement": statement, "person": request.user.person, "queue1": None, "queue2": None, "categ": None})        
    # № очереди
    queue1 = get_queue(request.user.person.id, 0)
    c=0
    for s in statement:
       c = s.category_id
    queue2 = get_queue(request.user.person.id, c)
    # Категория
    s = Statement.objects.get(person_id=request.user.person.id)
    categ = s.category
    return render(request, "statement/list.html", {"statement": statement, "person": request.user.person, "queue1": queue1, "queue2": queue2, "categ": categ})

def export_word(request): 
    # Create a HttpResponse object and set its content_type header value to Microsoft word.
    response = HttpResponse(content_type='application/vnd.ms-word') 
    # Set HTTP response Content-Disposition header value. Tell web server client the attached file name is students.xls.
    response['Content-Disposition'] = 'attachment;filename=statement.doc' 

    statement = Statement.objects.get(person_id=request.user.person.id)
    queue = Queue.objects.all().filter(dateq2__isnull=True)

    document = Document()

    document.add_heading('Заявление', 0)
    p = document.add_paragraph('Электронная очередь “Uya”')
    p = document.add_paragraph(str(statement.dates.strftime('%d.%m.%Y')))
    p = document.add_paragraph(str(statement.category))
    document.add_heading(str(statement.person), level=1)

    if statement.examination is not None:
        document.add_heading('Проверка документов', level=2)
        records = (
            (str(statement.datee.strftime('%d.%m.%Y %H:%M')), str(statement.examination.department)),
        )
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Дата'
        hdr_cells[1].text = 'Отдел'
        for date, department in records:
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = department
        
    if statement.sign1 is not None:
        document.add_heading('Подписание заявления', level=2)
        records = (
            (str(statement.dates1.strftime('%d.%m.%Y %H:%M')), str(statement.sign1.department)),
            (str(statement.dates2.strftime('%d.%m.%Y %H:%M')), str(statement.sign2.department)),
            (str(statement.dates3.strftime('%d.%m.%Y %H:%M')), str(statement.sign3.department)),
            (str(statement.dates4.strftime('%d.%m.%Y %H:%M')), str(statement.sign4.department)),
            (str(statement.dates5.strftime('%d.%m.%Y %H:%M')), str(statement.sign5.department)),
        )
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Дата'
        hdr_cells[1].text = 'Отдел'
        for date, department in records:
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = department

    document.add_heading('Статус', level=2)
    p = document.add_paragraph(str(statement.status)).bold = True

    document.add_heading('№ в общей очереди', level=2)
    p = document.add_paragraph(str(get_queue(request.user.person.id, 0))).bold = True

    document.add_heading('№ в очереди по своей категории', level=2)
    p = document.add_paragraph(str(get_queue(request.user.person.id, statement.category_id))).bold = True

    #document.add_page_break()

    # Create a StringIO object.
    output = BytesIO()
    # Save the workbook data to the above StringIO object.
    document.save(output)
    # Reposition to the beginning of the StringIO object.
    output.seek(0)
    # Write the StringIO object's value to HTTP response to send the word file to the web server client.
    response.write(output.getvalue()) 
    return response

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
def statement_create(request):
    try:

        if request.method == "POST":
            statement = Statement()        
            #statement.dates = request.POST.get("dates")
            statement.dates = datetime.datetime.now()
            #statement.person = request.POST.get("person")        
            statement.person_id = request.user.person.id
            statement.category = Category.objects.filter(id=request.POST.get("category")).first()
            statement.status_id = 1
            if 'paper1' in request.FILES:                
                statement.kind1 = Kind.objects.filter(id=request.POST.get("kind1")).first()
                statement.paper1 = request.FILES['paper1']        
            if 'paper2' in request.FILES:                
                statement.kind2 = Kind.objects.filter(id=request.POST.get("kind2")).first()
                statement.paper2 = request.FILES['paper2']        
            if 'paper3' in request.FILES:                
                statement.kind3 = Kind.objects.filter(id=request.POST.get("kind3")).first()
                statement.paper3 = request.FILES['paper3']        
            if 'paper4' in request.FILES:                
                statement.kind4 = Kind.objects.filter(id=request.POST.get("kind4")).first()
                statement.paper4 = request.FILES['paper4']        
            if 'paper5' in request.FILES:                
                statement.kind5 = Kind.objects.filter(id=request.POST.get("kind5")).first()
                statement.paper5 = request.FILES['paper5']        
            if 'paper6' in request.FILES:                
                statement.kind6 = Kind.objects.filter(id=request.POST.get("kind6")).first()
                statement.paper6 = request.FILES['paper6']        
            if 'paper7' in request.FILES:                
                statement.kind7 = Kind.objects.filter(id=request.POST.get("kind7")).first()
                statement.paper7 = request.FILES['paper7']        
            statement.save()
            # Записать лог
            statement = Statement.objects.all().order_by("-id")[0]
            print(statement)
            print(statement.id)
            history_create(request, request.user.person, statement, "Заявка подана", None)        
            return HttpResponseRedirect(reverse('statement_list'))
        else:
            print(request.user.person)
            statementform = StatementCreateForm(initial={'dates': datetime.datetime.now().strftime('%Y-%m-%d'), 'person': request.user.person })
            return render(request, "statement/create.html", {"form": statementform})
    except Exception as exception:
        print(exception)
        return HttpResponse(exception)

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def statement_check(request, id):
    try:
        statement = Statement.objects.get(id=id) 
        if request.method == "POST":
            statement.examination_id = request.user.employee.id
            #statement.datee = request.POST.get("datee")
            statement.datee = datetime.datetime.now()
            message=""
            # Определить какая кнопка нажата
            if 'btnAccept' in request.POST:
                statement.status_id = 3
                message="Документы приняты"
                #print("Accept")
            if 'btnReturn' in request.POST:
                message="Документы возвращены. " + request.POST.get("Comment")            
                statement.status_id = 2
                #print("Return")
            statement.save()
            # Записать лог
            #print(statement)
            #print(statement.id)
            history_create(request, statement.person, statement, message, request.user.employee )  
            return HttpResponseRedirect(reverse('statement_index'))
        else:
            # Загрузка начальных данных
            statementform = StatementCheckForm(initial={'dates': statement.dates.strftime('%Y-%m-%d'), 'person': statement.person, 'category': statement.category, 'examination': statement.examination, 'datee': datetime.datetime.now().strftime('%Y-%m-%d'),  })
            return render(request, "statement/check.html", {"form": statementform, "statement": statement})
    except Statement.DoesNotExist:
        return HttpResponseNotFound("<h2>Statement not found</h2>")

# Функция edit выполняет редактирование объекта.
@login_required
def statement_edit(request, id):
    try:
        statement = Statement.objects.get(id=id) 
        if request.method == "POST":
            #statement.dates = request.POST.get("dates")
            statement.dates = datetime.datetime.now()
            #statement.person = request.POST.get("person")        
            statement.person_id = request.user.person.id
            statement.category = Category.objects.filter(id=request.POST.get("category")).first()
            statement.status_id = 1
            if 'paper1' in request.FILES:                
                statement.kind1 = Kind.objects.filter(id=request.POST.get("kind1")).first()
                statement.paper1 = request.FILES['paper1']        
            if 'paper2' in request.FILES:                
                statement.kind2 = Kind.objects.filter(id=request.POST.get("kind2")).first()
                statement.paper2 = request.FILES['paper2']        
            if 'paper3' in request.FILES:                
                statement.kind3 = Kind.objects.filter(id=request.POST.get("kind3")).first()
                statement.paper3 = request.FILES['paper3']        
            if 'paper4' in request.FILES:                
                statement.kind4 = Kind.objects.filter(id=request.POST.get("kind4")).first()
                statement.paper4 = request.FILES['paper4']        
            if 'paper5' in request.FILES:                
                statement.kind5 = Kind.objects.filter(id=request.POST.get("kind5")).first()
                statement.paper5 = request.FILES['paper5']        
            if 'paper6' in request.FILES:                
                statement.kind6 = Kind.objects.filter(id=request.POST.get("kind6")).first()
                statement.paper6 = request.FILES['paper6']        
            if 'paper7' in request.FILES:                
                statement.kind7 = Kind.objects.filter(id=request.POST.get("kind7")).first()
                statement.paper7 = request.FILES['paper7']        
            statement.save()
            # Записать лог
            statement = Statement.objects.all().order_by("-id")[0]
            #print(statement)
            #print(statement.id)
            history_create(request, request.user.person, statement, "Заявка изменен", None)    
            return HttpResponseRedirect(reverse('statement_list'))
        else:
            # Загрузка начальных данных
            statementform = StatementEditForm(initial={'dates': statement.dates.strftime('%Y-%m-%d'), 'person': statement.person, 'category': statement.category, 'kind1': statement.kind1, 'paper1': statement.paper1, 'kind2': statement.kind2, 'paper2': statement.paper2, 'kind3': statement.kind3, 'paper3': statement.paper3, 'kind4': statement.kind4, 'paper4': statement.paper4, 'kind5': statement.kind5, 'paper5': statement.paper5, 'kind6': statement.kind6, 'paper6': statement.paper6, 'kind7': statement.kind7, 'paper7': statement.paper7, })
            return render(request, "statement/edit.html", {"form": statementform})
    except Statement.DoesNotExist:
        return HttpResponseNotFound("<h2>Statement not found</h2>")

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Commissions")
def statement_subscribe(request, id):
    try:
        #print('employee', request.user.employee)
        #print('id', request.user.employee.id)
        #print('signature', request.user.employee.signature)                
        statement = Statement.objects.get(id=id)
        if request.method == "POST":
            if (request.user.employee.signature==1):
                statement.sign1_id = request.user.employee.id
                statement.dates1 = datetime.datetime.now()
                history_create(request, statement.person, statement, str(request.user.employee.department) + ": Заявление одобрено", request.user.employee )          
            if (request.user.employee.signature==2):
                statement.sign2_id = request.user.employee.id
                statement.dates2 = datetime.datetime.now()
                history_create(request, statement.person, statement, str(request.user.employee.department) + ": Заявление одобрено", request.user.employee )          
            if (request.user.employee.signature==3):
                statement.sign3_id = request.user.employee.id
                statement.dates3 = datetime.datetime.now()
                history_create(request, statement.person, statement, str(request.user.employee.department) + ": Заявление одобрено", request.user.employee )          
            if (request.user.employee.signature==4):
                statement.sign4_id = request.user.employee.id
                statement.dates4 = datetime.datetime.now()
                history_create(request, statement.person, statement, str(request.user.employee.department) + ": Заявление одобрено", request.user.employee )          
            if (request.user.employee.signature==5):
                statement.status_id = 4
                statement.sign5_id = request.user.employee.id
                statement.dates5 = datetime.datetime.now()
                history_create(request, statement.person, statement, str(request.user.employee.department) + ": Постановка в очередь", request.user.employee )
                queue_create(request, statement, "Постановка в очередь")          
            statement.save()            
            return HttpResponseRedirect(reverse('statement_index'))
        else:
            # Загрузка начальных данных
            statementform = StatementSubscribeForm(initial={'dates': statement.dates.strftime('%Y-%m-%d'), 'person': statement.person, 'category': statement.category, 'sign1': statement.sign1, 'sign2': statement.sign2, 'sign3': statement.sign3, 'sign': statement.sign4, 'sign5': statement.sign5, })
            return render(request, "statement/subscribe.html", {"form": statementform, "statement": statement})
    except Statement.DoesNotExist:
        return HttpResponseNotFound("<h2>Statement not found</h2>")
    
# Просмотр страницы read.html для просмотра объекта.
@login_required
def statement_read(request, id):
    try:
        statement = Statement.objects.get(id=id) 
        return render(request, "statement/read.html", {"statement": statement})
    except Request.DoesNotExist:
        return HttpResponseNotFound("<h2>Statement not found</h2>")

# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers", "Commissions", "Managements")
def queue_index(request):
    queue = Queue.objects.all().order_by('dateq')
    #print('signature', request.user.employee.signature) 
    return render(request, "queue/index.html", {"queue": queue})
    
# Список 
@login_required
def queue_list(request):
    #queue = Queue.objects.all().order_by('dateq')
    print(request.user.person.id)
    queue = Queue.objects.filter(person_id=request.user.person.id)
    return render(request, "queue/list.html", {"queue": queue, "person": request.user.person})

# Запись в очередь
@login_required
def queue_create(request, _statement, _details ):
    queue = Queue()        
    queue.statement = _statement
    #queue.dateq = dateq
    queue.details = _details
    queue.save()

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def queue_edit(request, id):
    try:
        queue = Queue.objects.get(id=id)
        if queue.dateq2 is None:
            queue.dateq2 = datetime.datetime.now()
        if request.method == "POST":
            queue.dateq2 = request.POST.get("dateq2")
            queue.details2 = request.POST.get("details2")
            queue.save()
            return HttpResponseRedirect(reverse('queue_index'))
        else:
            # Загрузка начальных данных
            queueform = QueueEditForm(initial={'dateq2': queue.dateq2.strftime('%Y-%m-%d'),'details2': queue.details2,})
            return render(request, "queue/edit.html", {"form": queueform, 'dates': queue.statement.dates, 'person': queue.statement.person, 'category': queue.statement.category, 'status': queue.statement.status, 'dateq': queue.dateq, 'details': queue.details})
    except Queue.DoesNotExist:
        return HttpResponseNotFound("<h2>Queue not found</h2>")
    
# Просмотр страницы read.html для просмотра объекта.
@login_required
def queue_read(request, id):
    try:
        queue = Queue.objects.get(id=id) 
        return render(request, "queue/read.html", {"queue": queue})
    except Request.DoesNotExist:
        return HttpResponseNotFound("<h2>Queue not found</h2>")

# Полоучить порядковый номер в очереди клиент person_id категория category_id
def get_queue(person_id,category_id):
    qq = 0
    # Общая очередь
    if category_id == 0:
        queue = Queue.objects.filter(dateq2__isnull=True).order_by('dateq')
    # Или по выбранной категории
    else:
        # используется подазпрос
        statement_query = Statement.objects.filter(category_id=category_id).only('id').all()
        queue = Queue.objects.filter(dateq2__isnull=True).filter(id__in=statement_query).order_by('dateq')
    x=0
    for q in queue:
        x = x+1
        if q.statement.person_id == person_id:
            qq = x
    return qq
    
# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def kind_index(request):
    kind = Kind.objects.all().order_by('title')
    return render(request, "kind/index.html", {"kind": kind})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def kind_create(request):
    if request.method == "POST":
        kind = Kind()        
        kind.title = request.POST.get("title")
        kind.save()
        return HttpResponseRedirect(reverse('kind_index'))
    else:        
        kindform = KindForm(request.FILES)
        return render(request, "kind/create.html", {"form": kindform})

# Функция edit выполняет редактирование объекта.
@login_required
@group_required("Managers")
def kind_edit(request, id):
    try:
        kind = Kind.objects.get(id=id) 
        if request.method == "POST":
            kind.title = request.POST.get("title")
            kind.save()
            return HttpResponseRedirect(reverse('kind_index'))
        else:
            # Загрузка начальных данных
            kindform = KindForm(initial={'title': kind.title, })
            return render(request, "kind/edit.html", {"form": kindform})
    except Kind.DoesNotExist:
        return HttpResponseNotFound("<h2>Kind not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def kind_delete(request, id):
    try:
        kind = Kind.objects.get(id=id)
        kind.delete()
        return HttpResponseRedirect(reverse('kind_index'))
    except Kind.DoesNotExist:
        return HttpResponseNotFound("<h2>Kind not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def kind_read(request, id):
    try:
        kind = Kind.objects.get(id=id) 
        return render(request, "kind/read.html", {"kind": kind})
    except Kind.DoesNotExist:
        return HttpResponseNotFound("<h2>Kind not found</h2>")    

# Список для изменения с кнопками создать, изменить, удалить
@login_required
def history_index(request):
    history = History.objects.all().order_by('-dateh')
    #print('signature', request.user.employee.signature) 
    return render(request, "history/index.html", {"history": history})

# Список 
@login_required
def history_list(request):
    #print(request.user.person.id)
    history = History.objects.filter(person_id=request.user.person.id).order_by('-dateh')
    return render(request, "history/list.html", {"history": history, "person": request.user.person})

# Запись в журнал
@login_required
def history_create(request, _person, _statement, _details, _employee ):
    history = History()        
    history.person = _person
    history.statement = _statement
    #history.dateh = _dateh
    history.details = _details
    history.employee = _employee
    history.save()
    
# Список для изменения с кнопками создать, изменить, удалить
@login_required
@group_required("Managers")
def news_index(request):
    #news = News.objects.all().order_by('surname', 'name', 'patronymic')
    #return render(request, "news/index.html", {"news": news})
    news = News.objects.all().order_by('-daten')
    return render(request, "news/index.html", {"news": news})

# Список для просмотра
def news_list(request):
    news = News.objects.all().order_by('-daten')
    return render(request, "news/list.html", {"news": news})

# В функции create() получаем данные из запроса типа POST, сохраняем данные с помощью метода save()
# и выполняем переадресацию на корень веб-сайта (то есть на функцию index).
@login_required
@group_required("Managers")
def news_create(request):
    if request.method == "POST":
        news = News()        
        news.daten = request.POST.get("daten")
        news.title = request.POST.get("title")
        news.details = request.POST.get("details")
        if 'photo' in request.FILES:                
            news.photo = request.FILES['photo']        
        news.save()
        return HttpResponseRedirect(reverse('news_index'))
    else:        
        #newsform = NewsForm(request.FILES, initial={'daten': datetime.datetime.now().strftime('%Y-%m-%d'),})
        newsform = NewsForm(initial={'daten': datetime.datetime.now().strftime('%Y-%m-%d'), })
        return render(request, "news/create.html", {"form": newsform})

# Функция edit выполняет редактирование объекта.
# Функция в качестве параметра принимает идентификатор объекта в базе данных.
@login_required
@group_required("Managers")
def news_edit(request, id):
    try:
        news = News.objects.get(id=id) 
        if request.method == "POST":
            news.daten = request.POST.get("daten")
            news.title = request.POST.get("title")
            news.details = request.POST.get("details")
            if "photo" in request.FILES:                
                news.photo = request.FILES["photo"]
            news.save()
            return HttpResponseRedirect(reverse('news_index'))
        else:
            # Загрузка начальных данных
            newsform = NewsForm(initial={'daten': news.daten.strftime('%Y-%m-%d'), 'title': news.title, 'details': news.details, 'photo': news.photo })
            return render(request, "news/edit.html", {"form": newsform})
    except News.DoesNotExist:
        return HttpResponseNotFound("<h2>News not found</h2>")

# Удаление данных из бд
# Функция delete аналогичным функции edit образом находит объет и выполняет его удаление.
@login_required
@group_required("Managers")
def news_delete(request, id):
    try:
        news = News.objects.get(id=id)
        news.delete()
        return HttpResponseRedirect(reverse('news_index'))
    except News.DoesNotExist:
        return HttpResponseNotFound("<h2>News not found</h2>")

# Просмотр страницы read.html для просмотра объекта.
@login_required
def news_read(request, id):
    try:
        news = News.objects.get(id=id) 
        return render(request, "news/read.html", {"news": news})
    except News.DoesNotExist:
        return HttpResponseNotFound("<h2>News not found</h2>")

# Регистрационная форма 
def signup(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            auth_login(request, user)
            #print(user.id)
            # Добавить профиль
            person = Person.objects.all().order_by("-id")[0]
            id = person.id + 1            
            person = Person()
            person.id = id
            person.user = user
            person.sex = "М"
            person.iin = str(random.randint(10000000000,999999999999))
            person.surname = "Фамилия"
            person.name = "Имя"
            person.patronymic = "Отчество"
            person.birthday = datetime.datetime.now()
            person.save()
            #print(person)
            #print(person.id)            
            return redirect('person_edit')
            #return redirect('index')
            #return render(request, 'registration/register_done.html', {'new_user': user})
    else:
        form = SignUpForm()
    return render(request, 'registration/signup.html', {'form': form})


@method_decorator(login_required, name='dispatch')
class UserUpdateView(UpdateView):
    model = User
    #fields = ('first_name', 'last_name', 'email',)
    fields = ('email',)
    template_name = 'registration/my_account.html'
    success_url = reverse_lazy('index')
    #success_url = reverse_lazy('my_account')
    def get_object(self):
        return self.request.user

# Выход
from django.contrib.auth import logout
def logoutUser(request):
    logout(request)
    return render(request, "index.html")

